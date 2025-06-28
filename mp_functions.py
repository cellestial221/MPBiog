from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn  # Changed from ns to qn
from datetime import datetime
from bs4 import BeautifulSoup
import docx.opc.constants
import anthropic
import os
import PyPDF2
import re
import requests
import io
import wikipediaapi

def verify_constituency_in_wikipedia(page_url, constituency):
    """
    Simple verification: check if constituency appears in Wikipedia page content
    Returns True if constituency is mentioned, False otherwise
    """
    try:
        import requests
        from bs4 import BeautifulSoup

        response = requests.get(page_url, timeout=10)
        if response.status_code != 200:
            return False

        soup = BeautifulSoup(response.text, 'html.parser')
        content = soup.get_text().lower()

        # Check for constituency (handle common variations)
        constituency_lower = constituency.lower()
        constituency_variations = [
            constituency_lower,
            constituency_lower.replace(' and ', ' & '),
            constituency_lower.replace(' & ', ' and '),
        ]

        for variation in constituency_variations:
            if variation in content:
                return True

        return False

    except Exception as e:
        print(f"Error verifying constituency: {str(e)}")
        return False


def get_mp_wiki_link_verified(mp_name, constituency):
    """
    Enhanced version of get_mp_wiki_link that verifies constituency
    - Uses the same search logic as before
    - Adds constituency verification for found pages
    - Future-proof: doesn't rely on specific election years
    """
    try:
        import requests
        from bs4 import BeautifulSoup

        print(f"Searching for Wikipedia page: {mp_name} (constituency: {constituency})")

        # Use generic "List of MPs" page (not year-specific)
        potential_matches = []

        # Try current parliament list first
        list_urls = [
            "https://en.wikipedia.org/wiki/List_of_MPs_elected_in_the_2024_United_Kingdom_general_election",
            "https://en.wikipedia.org/wiki/List_of_current_members_of_the_British_Parliament"
        ]

        name_parts = mp_name.lower().split()
        first_name = name_parts[0] if name_parts else ""
        last_name = name_parts[-1] if len(name_parts) > 1 else name_parts[0] if name_parts else ""

        for list_url in list_urls:
            try:
                response = requests.get(list_url, timeout=10)
                if response.status_code != 200:
                    continue

                soup = BeautifulSoup(response.text, 'html.parser')
                tables = soup.find_all('table', class_='wikitable')

                for table in tables:
                    for link in table.find_all('a'):
                        href = link.get('href', '')
                        if not href.startswith('/wiki/'):
                            continue

                        # Skip obvious non-MP pages
                        if any(x in href.lower() for x in [
                            'constituency', 'party', 'election', 'parliament', 'list_of', 'category:'
                        ]):
                            continue

                        link_text = link.get_text().strip().lower()

                        # Simple name matching (same as before)
                        if (mp_name.lower() in link_text or
                            (first_name in link_text and last_name in link_text) or
                            (last_name in link_text and len(last_name) > 4)):

                            full_url = f"https://en.wikipedia.org{href}"
                            if full_url not in [m['url'] for m in potential_matches]:
                                potential_matches.append({
                                    'url': full_url,
                                    'text': link.get_text().strip()
                                })

            except Exception as e:
                print(f"Error searching {list_url}: {str(e)}")
                continue

        print(f"Found {len(potential_matches)} potential Wikipedia matches")

        # Now verify each match has the constituency
        for match in potential_matches:
            print(f"Checking: {match['text']} -> {match['url']}")

            if verify_constituency_in_wikipedia(match['url'], constituency):
                print(f"✅ VERIFIED: Constituency '{constituency}' found in {match['text']}")
                return match['url']
            else:
                print(f"❌ REJECTED: Constituency '{constituency}' not found in {match['text']}")

        print(f"❌ No verified Wikipedia page found for {mp_name} in {constituency}")
        return None

    except Exception as e:
        print(f"Error in Wikipedia search: {str(e)}")
        return None


def get_wiki_data_verified(mp_name, constituency, max_chars=3500):
    """
    Enhanced version of get_wiki_data with constituency verification
    - Same content extraction logic as before
    - Uses verified Wikipedia URL
    """
    try:
        print(f"Getting verified Wikipedia data for: {mp_name}")

        # Get verified Wikipedia URL
        wiki_url = get_mp_wiki_link_verified(mp_name, constituency)

        if not wiki_url:
            print("No verified Wikipedia page found")
            return None

        # Extract page title from URL
        page_title = wiki_url.split('/')[-1].replace('_', ' ')
        print(f"Using verified Wikipedia page: {page_title}")

        # Get content using Wikipedia API (same as original function)
        import wikipediaapi

        wiki = wikipediaapi.Wikipedia(
            user_agent='MP_Biography_Generator (yourname@example.com)',
            language='en'
        )

        page = wiki.page(page_title)

        if not page.exists():
            print(f"Page does not exist in Wikipedia API: {page_title}")
            return None

        # Get content (exactly same logic as original get_wiki_data function)
        content = page.summary
        print(f"Summary length: {len(content)} characters")

        # Add important sections (same as original)
        important_sections = [
            'early life', 'education', 'background', 'career', 'personal life',
            'political career', 'parliamentary career', 'political views',
            'controversies', 'awards', 'publications'
        ]

        essential_added = False
        for section in page.sections:
            section_title_lower = section.title.lower()

            if any(x in section_title_lower for x in ['see also', 'references', 'external links', 'notes', 'bibliography']):
                continue

            is_important = any(imp in section_title_lower for imp in important_sections)

            if is_important:
                section_text = f"\n\n{section.title}\n{section.text}"
                if len(content) + len(section_text) <= max_chars:
                    content += section_text
                    essential_added = True
                    print(f"Added section: {section.title}")

        # Add other sections if room (same logic as original)
        if not essential_added or len(content) < max_chars * 0.7:
            for section in page.sections:
                section_title_lower = section.title.lower()

                if (any(imp in section_title_lower for imp in important_sections) or
                    any(x in section_title_lower for x in ['see also', 'references', 'external links', 'notes', 'bibliography'])):
                    continue

                section_text = f"\n\n{section.title}\n{section.text}"
                if len(content) + len(section_text) <= max_chars:
                    content += section_text

        print(f"Final verified content length: {len(content)} characters")
        return content

    except Exception as e:
        print(f"Error in verified Wikipedia data retrieval: {str(e)}")
        return None


def get_wiki_url_verified(mp_name, constituency):
    """
    Simple wrapper to get verified Wikipedia URL
    """
    return get_mp_wiki_link_verified(mp_name, constituency)


def search_perplexity(mp_name, issues, api_key):
    """
    Search for MP's statements on specific issues using Perplexity AI API

    Args:
        mp_name (str): The name of the MP
        issues (str): Issues to search for
        api_key (str): Perplexity API key

    Returns:
        str: Search results formatted for inclusion in the biography
    """
    import requests
    import json

    # Prepare the API request
    url = "https://api.perplexity.ai/chat/completions"

    # Format the query to search specifically on parliament.uk
    query = f"Has {mp_name} MP ever said anything or done about {issues}? Find specific quotes, statements, actions, or positions they have taken on this issue. Focus on parliamentary records, speeches, votes, or committee work. Be specific and accurate and concise, list out the remarks one by one. If no relevant remarks are found, please reply succinctly confirming so - no need for a long explanation. site:parliament.uk"

    # Prepare the payload
    payload = {
        "model": "sonar-deep-research",
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful research assistant. Search for accurate information about UK Members of Parliament and their positions on issues. Provide detailed citations to parliamentary records where possible."
            },
            {
                "role": "user",
                "content": query
            }
        ],
        "max_tokens": 1000,
        "temperature": 0.2,
        "search_domain_filter": None,
        "return_images": False,
        "return_related_questions": False,
        "stream": False
    }

    # Set up headers with the API key
    headers = {
        "accept": "application/json",
        "content-type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }

    # Make the API request
    try:
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()  # Raise exception for HTTP errors

        # Parse the response
        result = response.json()

        # Extract the assistant's message content
        if "choices" in result and len(result["choices"]) > 0:
            content = result["choices"][0]["message"]["content"]
            citations = result.get("citations", [])

            # Format the content for inclusion in the biography
            formatted_result = content

            return formatted_result
        else:
            return f"No information found about {mp_name}'s statements on {issues}."

    except Exception as e:
        print(f"Error with Perplexity API: {str(e)}")
        return f"Error searching for information: {str(e)}"


def create_hyperlink(paragraph, text, url):
    """Create a hyperlink in a paragraph"""
    # Create relationship
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create new run
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set font
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'), 'Hanken Grotesk')
    rFont.set(qn('w:hAnsi'), 'Hanken Grotesk')
    rPr.append(rFont)

    # Set font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)

    # Set color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')  # Blue
    rPr.append(color)

    # Add underline for links
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Add text
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def clean_text(text):
    """Clean up the text format"""
    # Handle list or string input
    if isinstance(text, list):
        text = '\n'.join(str(item) for item in text)

    # Convert to string if it isn't already
    text = str(text)

    # Remove TextBlock wrapper
    if 'TextBlock(text=' in text:
        text = text.split('TextBlock(text=\'')[1].split('\', type=\'text\')')[0]

    # Replace escaped characters
    text = text.replace('\\n', '\n')
    text = text.replace('\\\'', "'")

    # Replace straight quotes with curly quotes
    text = text.replace('"', '"').replace('"', '"')

    # Remove empty lines while preserving intentional spacing
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        if line.strip():
            cleaned_lines.append(line.strip())
        else:
            cleaned_lines.append('')  # Preserve empty lines for spacing

    return '\n'.join(cleaned_lines)

def get_mp_id(mp_name):
    """Search for MP and return their ID"""
    try:
        # Make API request to search for MP
        search_url = f"https://members-api.parliament.uk/api/Members/Search?Name={mp_name}"
        response = requests.get(search_url)
        if response.status_code == 200:
            data = response.json()
            if data['items'] and len(data['items']) > 0:
                return data['items'][0]['value']['id']
    except Exception as e:
        print(f"Error searching for MP: {str(e)}")
    return None

def get_verified_positions(mp_id):
    """Get verified data from Parliament API with debug output"""
    print(f"DEBUG: Starting get_verified_positions for MP ID: {mp_id}")

    try:
        verified_data = {
            'current_committees': [],
            'historical_committees': [],
            'current_roles': [],
            'historical_roles': [],
            'synopsis': None,
            'recent_contributions': None,
            'api_response': None
        }

        if not mp_id:
            print("DEBUG: No MP ID provided")
            return verified_data

        print("DEBUG: About to fetch synopsis...")
        # Get synopsis
        synopsis_url = f"https://members-api.parliament.uk/api/Members/{mp_id}/Synopsis"
        print(f"DEBUG: Calling synopsis URL: {synopsis_url}")
        synopsis_response = requests.get(synopsis_url, timeout=10)
        print(f"DEBUG: Synopsis response: {synopsis_response.status_code}")

        if synopsis_response.status_code == 200:
            synopsis_data = synopsis_response.json()
            if 'value' in synopsis_data:
                verified_data['synopsis'] = synopsis_data['value']
                print(f"DEBUG: Synopsis found: {len(verified_data['synopsis'])} characters")

        print("DEBUG: Synopsis completed, moving to contributions...")

        # Get contribution summary
        print("DEBUG: About to fetch contribution summary...")
        contributions_url = f"https://members-api.parliament.uk/api/Members/{mp_id}/ContributionSummary"
        print(f"DEBUG: Calling contributions URL: {contributions_url}")
        contributions_response = requests.get(contributions_url, timeout=10)
        print(f"DEBUG: Contributions response: {contributions_response.status_code}")

        if contributions_response.status_code == 200:
            contributions_data = contributions_response.json()
            if 'items' in contributions_data and contributions_data['items']:
                # Get most recent contributions (last 30 days)
                recent_contributions = {
                    'total_count': 0,
                    'recent_debates': []
                }

                for item in contributions_data['items'][:5]:  # Look at 5 most recent
                    contribution = item['value']
                    recent_contributions['total_count'] += contribution.get('totalContributions', 0)
                    recent_contributions['recent_debates'].append({
                        'title': contribution.get('debateTitle'),
                        'date': contribution.get('sittingDate'),
                        'contributions': contribution.get('totalContributions', 0)
                    })

                verified_data['recent_contributions'] = recent_contributions
                print(f"DEBUG: Found {len(contributions_data['items'])} contribution items")

        print("DEBUG: Contributions completed, moving to biography...")

        # Get biography
        bio_url = f"https://members-api.parliament.uk/api/Members/{mp_id}/Biography"
        print(f"DEBUG: Calling biography URL: {bio_url}")
        bio_response = requests.get(bio_url, timeout=10)
        print(f"DEBUG: Biography response: {bio_response.status_code}")

        if bio_response.status_code == 200:
            print("DEBUG: Biography request successful, processing data...")
            bio_data = bio_response.json()['value']
            verified_data['api_response'] = bio_data
            print(f"DEBUG: Biography data keys: {list(bio_data.keys())}")

            # Process committee memberships - they're under 'committeeMemberships' directly
            if bio_data.get('committeeMemberships'):
                print(f"DEBUG: Found {len(bio_data['committeeMemberships'])} committee memberships")
                for committee in bio_data['committeeMemberships']:
                    committee_info = {
                        'name': committee.get('name'),
                        'start_date': committee.get('startDate', '')[:10] if committee.get('startDate') else None,
                        'end_date': committee.get('endDate', '')[:10] if committee.get('endDate') else 'present'
                    }

                    # Check if current or historical based on endDate
                    if not committee.get('endDate'):
                        verified_data['current_committees'].append(committee_info)
                    else:
                        verified_data['historical_committees'].append(committee_info)

            # Process government posts
            if bio_data.get('governmentPosts'):
                print(f"DEBUG: Found {len(bio_data['governmentPosts'])} government posts")
                for post in bio_data['governmentPosts']:
                    post_info = {
                        'name': post.get('name'),
                        'start_date': post.get('startDate', '')[:10] if post.get('startDate') else None,
                        'end_date': post.get('endDate', '')[:10] if post.get('endDate') else 'present'
                    }

                    if not post.get('endDate'):
                        verified_data['current_roles'].append(post_info)
                    else:
                        verified_data['historical_roles'].append(post_info)

            # Process opposition posts
            if bio_data.get('oppositionPosts'):
                print(f"DEBUG: Found {len(bio_data['oppositionPosts'])} opposition posts")
                for post in bio_data['oppositionPosts']:
                    post_info = {
                        'name': post.get('name'),
                        'start_date': post.get('startDate', '')[:10] if post.get('startDate') else None,
                        'end_date': post.get('endDate', '')[:10] if post.get('endDate') else 'present'
                    }

                    if not post.get('endDate'):
                        verified_data['current_roles'].append(post_info)
                    else:
                        verified_data['historical_roles'].append(post_info)

        print(f"DEBUG: Function completed successfully. Found {len(verified_data['current_committees'])} current committees, {len(verified_data['current_roles'])} current roles")
        return verified_data

    except Exception as e:
        print(f"DEBUG: Error in get_verified_positions: {str(e)}")
        import traceback
        print(f"DEBUG: Traceback: {traceback.format_exc()}")
        return verified_data

def get_mp_portrait(mp_id):
    """Get MP's thumbnail image"""
    if not mp_id:
        return None

    try:
        thumbnail_url = f"https://members-api.parliament.uk/api/Members/{mp_id}/Thumbnail"
        response = requests.get(thumbnail_url)
        if response.status_code == 200:
            return io.BytesIO(response.content)
    except Exception as e:
        print(f"Error getting MP thumbnail: {str(e)}")
    return None

def get_mp_data(mp_id):
    """Get comprehensive MP data from various API endpoints"""
    data = {
        'biography': None,
        'experience': None,
        'interests': None,
        'focus': None,
        'committees': None,
        'contact': None
    }

    if not mp_id:
        return data

    try:
        # Get biography
        bio_url = f"https://members-api.parliament.uk/api/Members/{mp_id}/Biography"
        bio_response = requests.get(bio_url)
        if bio_response.status_code == 200:
            data['biography'] = bio_response.json()['value']

        # Get experience
        exp_url = f"https://members-api.parliament.uk/api/Members/{mp_id}/Experience"
        exp_response = requests.get(exp_url)
        if exp_response.status_code == 200:
            data['experience'] = exp_response.json()['value']

        # Get registered interests
        interests_url = f"https://members-api.parliament.uk/api/Members/{mp_id}/RegisteredInterests"
        interests_response = requests.get(interests_url)
        if interests_response.status_code == 200:
            data['interests'] = interests_response.json()['value']

        # Get focus areas
        focus_url = f"https://members-api.parliament.uk/api/Members/{mp_id}/Focus"
        focus_response = requests.get(focus_url)
        if focus_response.status_code == 200:
            data['focus'] = focus_response.json()['value']

        # Get committee memberships from biography if available
        if data['biography']:
            data['committees'] = data['biography'].get('committeeMemberships', [])

        # Get contact information
        contact_url = f"https://members-api.parliament.uk/api/Members/{mp_id}/Contact"
        contact_response = requests.get(contact_url)
        if contact_response.status_code == 200:
            data['contact'] = contact_response.json()['value']

    except Exception as e:
        print(f"Error fetching MP data: {str(e)}")

    return data

def format_mp_data(mp_data):
    """Format MP data into structured text"""
    formatted_text = ""

    try:
        # Add biography information
        if mp_data['biography']:
            if mp_data['biography'].get('governmentPosts'):
                formatted_text += "\nGovernment Posts:\n"
                for post in mp_data['biography']['governmentPosts']:
                    formatted_text += f"- {post.get('name', '')}"
                    if post.get('startDate'):
                        formatted_text += f" ({post['startDate'][:10]} to {post.get('endDate', 'present')[:10]})\n"

            if mp_data['biography'].get('committeesMemberships'):
                formatted_text += "\nCommittee Memberships:\n"
                for committee in mp_data['biography']['committeeMemberships']:
                    formatted_text += f"- {committee.get('name', '')}"
                    if committee.get('startDate'):
                        formatted_text += f" ({committee['startDate'][:10]} to {committee.get('endDate', 'present')[:10]})\n"

        # Add focus areas
        if mp_data['focus']:
            formatted_text += "\nAreas of Focus:\n"
            for focus in mp_data['focus']:
                if focus.get('category'):
                    formatted_text += f"- {focus['category']}: "
                    formatted_text += ", ".join(focus.get('focus', [])) + "\n"

        # Add experience
        if mp_data['experience']:
            formatted_text += "\nProfessional Background:\n"
            for exp in mp_data['experience']:
                if exp.get('title'):
                    formatted_text += f"- {exp['title']}"
                    if exp.get('organisation'):
                        formatted_text += f", {exp['organisation']}"
                    formatted_text += "\n"

    except Exception as e:
        print(f"Error formatting MP data: {str(e)}")

    return formatted_text

def read_pdf(file_path):
    """Read and extract text from PDF"""
    try:
        with open(file_path, 'rb') as file:
            # Create PDF reader object
            pdf_reader = PyPDF2.PdfReader(file)

            # Extract text from all pages
            sections = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                # Clean up the text
                text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with single space
                sections.append(text)

            # Join all sections with proper spacing
            final_text = ' '.join(sections)

            # Extract key information
            summary_match = re.search(r'Summary(.*?)Current role', final_text, re.DOTALL)
            if summary_match:
                summary = summary_match.group(1).strip()
            else:
                summary = ""

            # Clean up and structure the text
            structured_text = "Summary:\n" + summary + "\n\n"

            # Add additional sections if found
            for section in ["Current role", "Previous roles", "World view and political affiliations"]:
                section_match = re.search(f"{section}(.*?)(Current|World view|Background|Election|\\Z)", final_text, re.DOTALL)
                if section_match:
                    structured_text += f"{section}:\n{section_match.group(1).strip()}\n\n"

            # Remove any registration dates, donor details, etc.
            structured_text = re.sub(r'\(Registered.*?\)', '', structured_text)
            structured_text = re.sub(r'Address of donor:.*?\n', '', structured_text)

            return structured_text.strip()

    except Exception as e:
        print(f"Error reading PDF: {str(e)}")
        return None

def read_example_bios():
    examples = []
    try:
        for file in os.listdir('example_bios'):
            if file.endswith('.docx'):
                print(f"Processing: {file}")
                doc = Document(os.path.join('example_bios', file))

                # Get all text content
                full_text = ' '.join(str(paragraph.text) for paragraph in doc.paragraphs)

                # Clean the text
                cleaned_text = clean_text(full_text)
                if cleaned_text:
                    examples.append(cleaned_text)
                print(f"Successfully processed {file}")
    except Exception as e:
        print(f"Error processing DOCX: {str(e)}")
        raise

    return '\n\n---\n\n'.join(examples)

def read_input_file(file_path):
    """Read and return the contents of the input file"""
    with open(file_path, 'r') as file:
        return file.read()

def find_mp_in_wikipedia_list(mp_name):
    """
    Find the MP in the list of elected MPs and return their Wikipedia page.

    Args:
        mp_name (str): The name of the MP to search for

    Returns:
        wikipediaapi.WikipediaPage or None: The Wikipedia page for the MP if found
    """
    try:
        wiki = wikipediaapi.Wikipedia(
            user_agent='MP_Biography_Generator (yourname@example.com)',
            language='en'
        )

        # Get the list of MPs page
        list_page = wiki.page("List_of_MPs_elected_in_the_2024_United_Kingdom_general_election")

        if not list_page.exists():
            print("MPs list page not found on Wikipedia")
            return None

        # Get all the links from the list page
        links = list_page.links

        # Normalize the MP name for matching (lowercase, no extra spaces)
        mp_name_norm = " ".join(mp_name.strip().lower().split())
        mp_surname = mp_name.strip().split()[-1].lower()

        # List to store potential matches
        potential_matches = []

        # Search through the links for potential MP matches
        for title, page in links.items():
            # Skip disambiguation, category, and list pages
            if any(x in title for x in ['List of', 'Category:', 'disambiguation', 'election']):
                continue

            # Skip pages about constituencies, parties, etc.
            if any(x in title.lower() for x in ['constituency', 'party', 'general election']):
                continue

            # Check if the title contains the MP's name
            if mp_name_norm in title.lower() or mp_surname in title.lower():
                potential_matches.append((title, page))

        print(f"Found {len(potential_matches)} potential matches for {mp_name}")

        # If we have potential matches, prioritize them by likelihood
        if potential_matches:
            # Check if any match contains "MP" or "politician"
            for title, page in potential_matches:
                if " MP" in title or "(politician)" in title:
                    print(f"Selected match with MP/politician indicator: {title}")
                    return page

            # If no matches with MP/politician indicators, use the first match
            print(f"Selected first potential match: {potential_matches[0][0]}")
            return potential_matches[0][1]

        return None

    except Exception as e:
        print(f"Error finding MP in Wikipedia list: {str(e)}")
        return None

def get_mp_wiki_link(mp_name):
    """
    Find the direct Wikipedia link for an MP from the list of elected MPs.

    Args:
        mp_name (str): The name of the MP to search for

    Returns:
        str or None: The Wikipedia URL if found, None otherwise
    """
    try:
        print(f"Searching for MP: {mp_name} in the 2024 elected MPs list")

        # Get the list page
        url = "https://en.wikipedia.org/wiki/List_of_MPs_elected_in_the_2024_United_Kingdom_general_election"
        response = requests.get(url)
        if response.status_code != 200:
            print(f"Error accessing MPs list page: Status code {response.status_code}")
            return None

        # Parse HTML
        soup = BeautifulSoup(response.text, 'html.parser')

        # Normalize the MP name for more flexible matching
        name_parts = mp_name.lower().split()
        first_name = name_parts[0] if name_parts else ""
        last_name = name_parts[-1] if len(name_parts) > 1 else name_parts[0] if name_parts else ""

        # Find all MP links in the tables - typically they're in wikitable class tables
        mp_links = []
        tables = soup.find_all('table', class_='wikitable')

        for table in tables:
            for link in table.find_all('a'):
                # Skip links with certain patterns
                if any(x in link.get('href', '').lower() for x in [
                    'constituency', 'party', 'election', 'parliament', 'list_of', 'category:'
                ]):
                    continue

                link_text = link.get_text().strip().lower()

                # Different matching strategies
                if (mp_name.lower() in link_text or                           # Full name match
                    (first_name in link_text and last_name in link_text) or   # First and last name in text
                    last_name in link_text):                                  # Last name at minimum

                    href = link.get('href')
                    if href and href.startswith('/wiki/'):
                        mp_links.append({
                            'text': link.get_text().strip(),
                            'url': f"https://en.wikipedia.org{href}",
                            'exact_match': mp_name.lower() == link_text
                        })

        # Print all potential matches for debugging
        print(f"Found {len(mp_links)} potential matches:")
        for i, link in enumerate(mp_links):
            print(f"{i+1}. {link['text']} - {link['url']}")

        # First try exact matches
        exact_matches = [link for link in mp_links if link['exact_match']]
        if exact_matches:
            print(f"Using exact match: {exact_matches[0]['text']} - {exact_matches[0]['url']}")
            return exact_matches[0]['url']

        # Then try best match (first match)
        if mp_links:
            print(f"Using best match: {mp_links[0]['text']} - {mp_links[0]['url']}")
            return mp_links[0]['url']

        print(f"No matches found for {mp_name}")
        return None

    except Exception as e:
        print(f"Error finding MP Wikipedia link: {str(e)}")
        return None


def get_wiki_data(mp_name, max_chars=3500):
    """Get comprehensive MP data from Wikipedia with length control"""
    try:
        print(f"\nAttempting to find Wikipedia data for: {mp_name}")

        # First try to get the direct link from the MPs list
        mp_wiki_url = get_mp_wiki_link(mp_name)

        if not mp_wiki_url:
            print("Could not find MP in the 2024 elected MPs list")
            return None

        # Extract page title from URL
        page_title = mp_wiki_url.split('/')[-1].replace('_', ' ')
        print(f"Found Wikipedia page: {page_title}")

        # Use the Wikipedia API to get full content
        wiki = wikipediaapi.Wikipedia(
            user_agent='MP_Biography_Generator (yourname@example.com)',
            language='en'
        )

        # Get the page using the API
        page = wiki.page(page_title)

        if not page.exists():
            print(f"Page does not exist in Wikipedia API: {page_title}")
            return None

        print(f"Successfully loaded page: {page.title}")

        # Get the full summary first
        content = page.summary
        print(f"Summary length: {len(content)} characters")

        # Define important biographical sections to prioritize
        important_sections = [
            'early life', 'education', 'background', 'career', 'personal life',
            'political career', 'parliamentary career', 'political views',
            'controversies', 'awards', 'publications'
        ]

        # Add essential sections first
        essential_added = False
        for section in page.sections:
            section_title_lower = section.title.lower()

            # Skip irrelevant sections
            if any(x in section_title_lower for x in ['see also', 'references', 'external links', 'notes', 'bibliography']):
                continue

            # Check if this is an important biographical section
            is_important = any(imp in section_title_lower for imp in important_sections)

            # Add important sections first
            if is_important:
                section_text = f"\n\n{section.title}\n{section.text}"

                # Check if adding this section would exceed max length
                if len(content) + len(section_text) <= max_chars:
                    content += section_text
                    essential_added = True
                    print(f"Added important section: {section.title} ({len(section_text)} chars)")

        # If we have room, add other sections that weren't important but might be relevant
        if not essential_added or len(content) < max_chars * 0.7:
            for section in page.sections:
                section_title_lower = section.title.lower()

                # Skip already processed important sections and irrelevant ones
                if (any(imp in section_title_lower for imp in important_sections) or
                    any(x in section_title_lower for x in ['see also', 'references', 'external links', 'notes', 'bibliography'])):
                    continue

                section_text = f"\n\n{section.title}\n{section.text}"

                # Check if adding this section would exceed max length
                if len(content) + len(section_text) <= max_chars:
                    content += section_text
                    print(f"Added additional section: {section.title} ({len(section_text)} chars)")

        # Add subsections if we still have room
        if len(content) < max_chars * 0.9:
            for section in page.sections:
                for subsection in section.sections:
                    subsection_title_lower = subsection.title.lower()

                    # Skip irrelevant subsections
                    if any(x in subsection_title_lower for x in ['see also', 'references', 'external links', 'notes']):
                        continue

                    subsection_text = f"\n\n{section.title} - {subsection.title}\n{subsection.text}"

                    # Check if adding this subsection would exceed max length
                    if len(content) + len(subsection_text) <= max_chars:
                        content += subsection_text
                        print(f"Added subsection: {subsection.title} ({len(subsection_text)} chars)")

        # Log summary of what we found
        print(f"Total content length: {len(content)} characters")
        print(f"Content sections: {content.count('\\n\\n')}")

        if len(content) < 500:
            print("WARNING: Retrieved content is very short. Content may be incomplete.")

        return content

    except Exception as e:
        print(f"Error fetching Wikipedia data: {str(e)}")
        return None

def get_wiki_url(mp_name):
    """Get Wikipedia URL for MP"""
    try:
        return get_mp_wiki_link(mp_name)
    except Exception as e:
        print(f"Error getting Wikipedia URL: {str(e)}")
        return None


# UPDATED GENERATE_BIOGRAPHY FUNCTION (mp_functions.py)
def generate_biography(mp_name, input_content, examples, verified_positions=None, comments=None, length_setting="medium"):
    # Validate and clean inputs (keep your existing logic)
    if isinstance(input_content, list):
        input_content = ' '.join(str(x) for x in input_content)
    if isinstance(examples, list):
        examples = ' '.join(str(x) for x in examples)

    input_content = str(input_content).strip()
    examples = str(examples).strip()

    if not input_content:
        input_content = f"Background information for {mp_name} could not be found. Further research is needed."

    # Get Wikipedia data as fallback
    wiki_content = get_wiki_data(mp_name)
    if wiki_content:
        input_content = f"{input_content}\n\nWikipedia information:\n{wiki_content}"

    client = anthropic.Client(api_key=os.getenv('ANTHROPIC_API_KEY'))
    current_date = datetime.now().strftime('%Y-%m-%d')

    # Create verified positions text (keep your existing logic)
    verified_positions_text = "\nVERIFIED PARLIAMENTARY INFORMATION:\n"
    if verified_positions:
        has_any_positions = False
        has_any_content = False

        if verified_positions.get('synopsis'):
            has_any_content = True
            verified_positions_text += f"\nOFFICIAL SYNOPSIS:\n{verified_positions['synopsis']}\n"

        if verified_positions['current_committees']:
            has_any_positions = True
            verified_positions_text += "\nCurrent Committee Memberships:\n"
            for committee in verified_positions['current_committees']:
                verified_positions_text += f"- {committee['name']} (Since {committee['start_date']})\n"

        if verified_positions['current_roles']:
            has_any_positions = True
            verified_positions_text += "\nCurrent Government/Opposition Roles:\n"
            for role in verified_positions['current_roles']:
                verified_positions_text += f"- {role['name']} (Since {role['start_date']})\n"

        if not has_any_positions:
            verified_positions_text += "\nNo current committee memberships or government/opposition roles found. Do not include any such positions in the biography.\n"
    else:
        verified_positions_text += "\nNo verified position data available. Do not include any committee memberships, government/opposition roles, or parliamentary activities in the biography.\n"

    # Create comments text (keep your existing logic)
    comments_text = ""
    if comments and len(comments) > 0:
        comments_text = "\n\nRELEVANT COMMENTS TO INCLUDE AT THE END OF THE BIOGRAPHY:\n"
        comments_text += "Please include a section at the end of the biography titled 'Relevant Comments'. "
        comments_text += "Format each comment as a bullet point (• ) item in a list. "
        comments_text += "Summarize each of these comments in a short paragraph, including the date. "
        comments_text += "Group similar comments together when appropriate. For dates, use British date format (day month year). "
        comments_text += "IMPORTANT: Process these comments in the exact order provided and include a reference marker [REF-X] at the end of each bullet point where X is the comment number (1, 2, 3, etc.).\n\n"

        for i, comment in enumerate(comments):
            comment_date = comment.get('date', '')
            try:
                if comment_date:
                    date_obj = datetime.strptime(comment_date, '%Y-%m-%d')
                    comment_date = date_obj.strftime('%d %B %Y')
            except:
                pass

            comments_text += f"Comment {i+1} [REF-{i+1}]:\n"
            comments_text += f"Type: {comment.get('type', '')}\n"
            comments_text += f"Date: {comment_date}\n"
            comments_text += f"URL: {comment.get('url', '')}\n"
            comments_text += f"Text: {comment.get('text', '')}\n\n"

    # Length-specific instructions (keep your existing logic)
    length_instructions = {
        "brief": {
            "description": "BRIEF biography (approximately 2-3 short paragraphs)",
            "structure": """Structure for BRIEF biography:
1. MP name and role as title
2. Party and constituency in parentheses
3. One paragraph introduction with current position and most important roles only
4. One paragraph covering the most significant career highlights only
5. One paragraph with key background information only""",
            "content_guidelines": """BRIEF content guidelines:
- Focus only on the most essential information
- Include current role and 1-2 most significant positions
- Mention only major career highlights (not detailed career history)
- Include only the most relevant educational background
- Omit minor roles, detailed chronology, and extensive background details
- Keep each paragraph to 2-3 sentences maximum
- Total length should be approximately 100-150 words"""
        },
        "medium": {
            "description": "STANDARD biography (the example length)",
            "structure": """Structure for STANDARD biography (follow the examples exactly):
1. MP name and role as title
2. Party and constituency in parentheses
3. Introduction paragraph with current position and verified roles
4. 'Politics' section with clear heading
5. 'Background' section with clear heading""",
            "content_guidelines": """STANDARD content guidelines:
- Follow the example biographies' length and detail level exactly
- Include comprehensive career history with specific positions and dates
- Provide detailed educational background and qualifications
- Include significant achievements and career progression
- Maintain the same level of detail as shown in the examples
- This is the default length that matches your training examples"""
        },
        "comprehensive": {
            "description": "COMPREHENSIVE biography (extended detail)",
            "structure": """Structure for COMPREHENSIVE biography:
1. MP name and role as title
2. Party and constituency in parentheses
3. Detailed introduction paragraph with current position and all verified roles
4. 'Politics' section with comprehensive political career details
5. 'Background' section with extensive career and education history
6. 'Early Life and Education' subsection if information available
7. 'Professional Career' subsection with detailed work history""",
            "content_guidelines": """COMPREHENSIVE content guidelines:
- Significantly expand on all sections compared to the examples
- Include detailed chronological career progression with specific dates
- Provide comprehensive educational background including institutions and qualifications
- Include extensive political career details, committee work, and parliamentary contributions
- Add more context about the significance of roles and achievements
- Include additional background information about early life if available
- Expand on professional experience with more company names, positions, and responsibilities
- Each section should be substantially longer than the standard examples
- Total length should be approximately 50-75% longer than the standard examples"""
        }
    }

    length_config = length_instructions.get(length_setting, length_instructions["medium"])

    # ===== PROMPT CACHING IMPLEMENTATION =====

    # Create the cached content (examples + instructions)
    cached_content = f"""Using these examples as a guide for style ONLY, generate a new biography for {mp_name}.

    LENGTH REQUIREMENT: {length_config['description']}

    {length_config['structure']}

    {length_config['content_guidelines']}

    CRITICAL REQUIREMENTS:
    1. ONLY include specific, verifiable facts about this MP - no generic statements about MP duties
    2. For career history, try to include named positions, companies, and years - no vague industry descriptions
    3. Focus on concrete facts: dates, organisations, position titles, qualifications, specific achievements
    4. DO NOT include any of these generic phrases or concepts:
       - "advocating for constituents' interests"
       - "matters that matter to constituents"

    Example biography for style reference (NOTE: This is STANDARD length):
    {examples}

    Important requirements:
    1. Match the exact formatting and style of the example, including the placement of newlines and section headers
    2. Use ONLY information from the provided input content and Wikipedia
    3. Rephrase and restructure the information - do not copy phrases directly
    4. Maintain the same professional tone and level of detail appropriate for the {length_setting} length
    5. Use clear section headers with proper spacing before and after - DO NOT use any markdown formatting like # for headers
    6. Focus on the most significant aspects of their career and current role, be specific and detailed
    7. Organise information chronologically within each section
    8. Keep sentences fairly concise, factual, and clear
    9. Use British English spelling AT ALL TIMES
    10. Do not include the detailed list of donations
    11. Include current significant roles and committee memberships in the top section
    12. Be VERY VERY careful in being accurate with dates, using today's current date ({current_date}) as reference to determine both past AND current roles
    13. Be sure to be VERY careful in being accurate with Committee names, memberships, and government roles if applicable
    14. If recent parliamentary contributions are provided, include a SHORT 1-2 sentence summary at the end of the Politics section
    15. Use the official synopsis where provided, incorporating its verified information naturally into the narrative
    16. Do NOT include Date of Birth
    17. In the Relevant Comments section, format each item as a bullet point (• ) and PRESERVE the reference markers [REF-X] at the end of each bullet point
    18. Do not repeat information given in prior sections, so make sure the information is in the relevant section and not elsewhere
    19. ADJUST THE TOTAL LENGTH according to the {length_setting} setting specified above"""

    # Create the dynamic content (MP-specific data)
    dynamic_content = f"""IMPORTANT: Use ONLY the following verified positions when mentioning committee memberships and roles.
    DO NOT list them explicitly, but incorporate them naturally into the narrative:
    {verified_positions_text}

    DO NOT include any committee memberships or roles that are not listed above, even if you find them in other sources.

    Information to use for the new biography:
    {input_content}

    {comments_text}"""

    # Structure messages with caching
    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": cached_content,
                    "cache_control": {"type": "ephemeral"}  # This caches the examples and instructions
                },
                {
                    "type": "text",
                    "text": dynamic_content  # This is unique per MP and not cached
                }
            ]
        }
    ]

    try:
        # Adjust max_tokens based on length setting
        max_tokens_map = {
            "brief": 1500,
            "medium": 3000,
            "comprehensive": 4500
        }

        response = client.messages.create(
            model="claude-sonnet-4-20250514",  # Updated to Sonnet 4
            max_tokens=max_tokens_map.get(length_setting, 3000),
            temperature=0.7,
            messages=messages
        )

        biography = str(response.content[0].text)
        return biography

    except Exception as e:
        print(f"Error in biography generation: {str(e)}")
        raise

# UPDATED MAIN FUNCTION (mp_functions.py) - Update the existing main() function
def main():
    # Check for API key
    if not os.getenv('ANTHROPIC_API_KEY'):
        print("Please set your ANTHROPIC_API_KEY first!")
        return

    # Get MP name
    mp_name = input("Enter MP name: ").strip()

    # Check if biography already exists
    if os.path.exists(f'new_bios/{mp_name}_biography.docx'):
        print(f"Biography already exists for {mp_name}")
        return

    # Get input file path
    input_path = input("Enter path to PDF file: ").strip()
    if not os.path.exists(input_path):
        print("File not found!")
        return

    try:
        # Read examples
        print("Reading example biographies...")
        examples = read_example_bios()

        # Read PDF input
        print("Reading PDF file...")
        input_content = read_pdf(input_path)
        has_pdf = input_content is not None

        # Get MP ID and API data
        mp_id = get_mp_id(mp_name)
        verified_positions = get_verified_positions(mp_id) if mp_id else None
        has_api_data = verified_positions is not None

        # Get Wikipedia data
        print("Fetching Wikipedia data...")
        wiki_data = get_wiki_data(mp_name)
        has_wiki_data = wiki_data is not None

        if input_content is None:
            print("Error reading PDF file")
            return

        # Generate biography with default medium length
        print("Generating biography...")
        biography = generate_biography(mp_name, input_content, examples, verified_positions, None, "medium")

        # Get Wikipedia URL if data exists
        wiki_url = get_wiki_url(mp_name) if has_wiki_data else None

        # Save biography
        print("Saving biography...")
        saved_path = save_biography(mp_name, biography,
                                has_pdf=has_pdf,
                                has_api_data=has_api_data,
                                has_wiki_data=has_wiki_data,
                                wiki_url=wiki_url)
        print(f"Biography saved to {saved_path}")

    except Exception as e:
        print(f"An error occurred: {str(e)}")


def apply_heading_style(paragraph):
    """Apply consistent styling to headings"""
    run = paragraph.runs[0]
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    paragraph.space_before = Pt(12)
    paragraph.space_after = Pt(6)

def apply_body_style(paragraph):
    """Apply consistent styling to body text"""
    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    paragraph.space_after = Pt(6)

# TEMPORARY DEBUG VERSION - Replace the save_biography function with this version
# This will print debug information to help identify the issue

# QUICK FIX: Add these lines at the VERY START of your save_biography function
# (Right after the function definition, before anything else)

def save_biography(mp_name, content, comments=None, has_pdf=False, has_api_data=False, has_wiki_data=False, wiki_url=None):
    """Save biography with hyperlinks - SIMPLE DEBUG VERSION"""

    # SIMPLE DEBUG - Add these lines at the start
    print(f"\n🚀 SAVE_BIOGRAPHY CALLED for {mp_name}")
    print(f"📄 Content length: {len(content) if content else 0}")
    print(f"💬 Comments: {comments}")
    print(f"🔗 Comments with URLs: {[c.get('url', 'NO_URL') for c in comments] if comments else 'NO_COMMENTS'}")

    # Check if content contains [REF-
    if content and '[REF-' in content:
        print(f"✅ Content contains REF markers!")
        import re
        markers = re.findall(r'\[REF-\d+\]', content)
        print(f"📝 Found markers: {markers}")
    else:
        print(f"❌ No REF markers found in content!")

    # Show a snippet of the Relevant Comments section
    if content and 'Relevant Comments' in content:
        lines = content.split('\n')
        in_comments = False
        print(f"📋 Relevant Comments section:")
        for line in lines:
            if line.strip() == 'Relevant Comments':
                in_comments = True
                continue
            if in_comments and line.strip():
                print(f"   {line}")
                if line.strip().startswith('•'):
                    break  # Just show first bullet point

    try:
        # Import at the start to catch any import errors
        import re
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from datetime import datetime

        print(f"✅ All imports successful")

        # Continue with your existing save_biography code here...
        # (Keep all the rest of your current save_biography function)

    except Exception as e:
        print(f"❌ ERROR in save_biography: {e}")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")
        raise

    import re

    doc = Document()

    # Define colors
    TEAL_COLOR = RGBColor(0, 161, 154)
    GREY_COLOR = RGBColor(128, 128, 128)

    # Add source information in header section (keeping existing logic)
    section = doc.sections[0]
    header = section.header
    source_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    # Get current date
    current_date = datetime.now().strftime('%d %B %Y')

    # Prepare source text parts
    source_text_start = f"AI generated MP biography on {current_date}. Sources: "
    sources = []
    if has_pdf:
        sources.append("user submitted PDF")
    if has_api_data:
        sources.append("Parliament's API data")

    # Add the start of the text
    run = source_para.add_run(source_text_start)
    run.font.name = 'Hanken Grotesk'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = GREY_COLOR

    if sources:
        run = source_para.add_run(" and ".join(sources))
        run.font.name = 'Hanken Grotesk'
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = GREY_COLOR

        if has_wiki_data and wiki_url:
            run = source_para.add_run(" and ")
            run.font.name = 'Hanken Grotesk'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = GREY_COLOR

            # Add Wikipedia as hyperlink
            create_hyperlink(source_para, "Wikipedia", wiki_url)
    else:
        run = source_para.add_run("none")
        run.font.name = 'Hanken Grotesk'
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = GREY_COLOR

    source_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Try to add MP's portrait
    mp_id = get_mp_id(mp_name)
    if mp_id:
        portrait = get_mp_portrait(mp_id)
        if portrait:
            doc.add_picture(portrait, width=Inches(2))
            doc.add_paragraph()

    # Clean and process the content
    content = clean_text(content)
    content = content.replace('# ', '')
    content = content.replace('## ', '')
    content = content.replace('#Politics', 'Politics')
    content = content.replace('#Background', 'Background')
    content = content.replace('#Relevant Comments', 'Relevant Comments')

    import re
    content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)

    # DEBUG: Print the comments data
    print(f"\n=== DEBUG SAVE_BIOGRAPHY ===")
    print(f"MP Name: {mp_name}")
    print(f"Comments received: {comments}")
    print(f"Number of comments: {len(comments) if comments else 0}")

    # IMPROVED COMMENT URL MAPPING
    comments_with_urls = []
    if comments:
        for i, comment in enumerate(comments):
            print(f"Comment {i}: {comment}")
            if comment.get('url'):
                comments_with_urls.append({
                    'index': i,
                    'url': comment['url'],
                    'date': comment.get('date', ''),
                    'text': comment.get('text', ''),
                    'type': comment.get('type', '')
                })
                print(f"  -> Added to comments_with_urls: {comment['url']}")

    print(f"DEBUG: Found {len(comments_with_urls)} comments with URLs")
    for i, comment in enumerate(comments_with_urls):
        print(f"  {i}: {comment['url']}")

    # COMPLETE FIXED VERSION - Replace the paragraph processing section

    # Replace the entire paragraph processing loop in your save_biography function with this:

    paragraphs = content.split('\n')
    in_comments_section = False
    bullet_point_count = 0

    # Process each paragraph
    for i in range(len(paragraphs)):
        para = paragraphs[i].strip()

        if para:
            p = doc.add_paragraph()

            # IMPROVED: More robust "Relevant Comments" detection
            if 'Relevant Comments' in para.strip():
                in_comments_section = True
                print("✅ DEBUG: Entered Relevant Comments section")
                print(f"   Line was: '{para.strip()}'")

            # Title (first paragraph)
            if i == 0:
                run = p.add_run(para.strip())
                run.font.name = 'Hanken Grotesk'
                run.font.size = Pt(10)
                run.font.weight = 700
                run.font.color.rgb = TEAL_COLOR

            # Party and constituency line
            elif para.strip().startswith('(Labour') or para.strip().startswith('(Conservative') or '(' in para.strip():
                run = p.add_run(para.strip())
                run.font.name = 'Hanken Grotesk'
                run.font.size = Pt(10)
                run.font.weight = 400
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Section headers
            elif (para.strip() in ['Politics', 'Background', 'Relevant Comments'] or
                  para.strip().endswith('Politics') or
                  para.strip().endswith('Background') or
                  para.strip().endswith('Relevant Comments')):

                # Clean the header text (remove any remaining #)
                header_text = para.strip().lstrip('#').strip()

                run = p.add_run(header_text)
                run.font.name = 'Hanken Grotesk'
                run.font.size = Pt(10)
                run.font.weight = 400  # Regular weight
                run.font.color.rgb = TEAL_COLOR  # Teal color like MP name

            # IMPROVED: More robust bullet point detection
            elif in_comments_section and (para.strip().startswith('•') or
                                         para.strip().startswith('-') or
                                         para.strip().startswith('*') or
                                         '[REF-' in para.strip()):

                print(f"\n🎯 FOUND BULLET POINT {bullet_point_count + 1}!")
                print(f"Original: '{para}'")
                print(f"First char: '{para.strip()[0] if para.strip() else 'EMPTY'}'")
                print(f"Contains REF: {'[REF-' in para}")

                # STRATEGY 1: Look for reference markers [REF-X]
                ref_match = re.search(r'\[REF-(\d+)\]', para)
                comment_url = None

                if ref_match:
                    ref_number = int(ref_match.group(1)) - 1  # Convert to 0-based index
                    print(f"✅ DEBUG: Found reference marker REF-{ref_number + 1}")

                    if ref_number < len(comments_with_urls):
                        comment_url = comments_with_urls[ref_number]['url']
                        print(f"✅ DEBUG: Matched by reference marker to URL: {comment_url}")

                        # Remove the reference marker from the text
                        para_cleaned = re.sub(r'\s*\[REF-\d+\]', '', para)
                        print(f"✅ DEBUG: Cleaned para: {para_cleaned}")
                        para = para_cleaned  # Update para with cleaned version
                    else:
                        print(f"❌ DEBUG: Reference number {ref_number + 1} exceeds available comments ({len(comments_with_urls)})")
                else:
                    print("⚠️ DEBUG: No reference marker found")

                # STRATEGY 2: Index-based matching (fallback)
                if not comment_url and comments_with_urls:
                    print(f"🔄 DEBUG: Trying index-based matching. Bullet count: {bullet_point_count}")
                    if bullet_point_count < len(comments_with_urls):
                        comment_url = comments_with_urls[bullet_point_count]['url']
                        print(f"✅ DEBUG: Matched by index {bullet_point_count} to URL: {comment_url}")

                # Add the main bullet point text
                # Add the main bullet point text - BUT handle punctuation
                bullet_text = para.strip()
                has_trailing_period = bullet_text.endswith('.')

                # Remove trailing period if present (we'll add it after the link)
                if has_trailing_period:
                    bullet_text = bullet_text.rstrip('.')

                print(f"📝 DEBUG: Adding text to document: {bullet_text}")
                run = p.add_run(bullet_text)


                run.font.name = 'Hanken Grotesk'
                run.font.size = Pt(10)
                run.font.weight = 300
                run.font.color.rgb = RGBColor(0, 0, 0)

                # Add the hyperlink if we found a URL
                if comment_url:
                    print(f"🔗 DEBUG: Adding hyperlink: {comment_url}")

                    # Add space and opening parenthesis
                    run = p.add_run(' (')
                    run.font.name = 'Hanken Grotesk'
                    run.font.size = Pt(10)
                    run.font.weight = 300
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    # Add the hyperlink
                    try:
                        create_hyperlink(p, 'link', comment_url)
                        print(f"✅ DEBUG: Successfully created hyperlink")
                    except Exception as e:
                        print(f"❌ DEBUG: Error creating hyperlink: {e}")

                    # Add closing parenthesis
                    run = p.add_run(')')
                    run.font.name = 'Hanken Grotesk'
                    run.font.size = Pt(10)
                    run.font.weight = 300
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    # Add period after the link if there was one originally
                    if has_trailing_period:
                        run = p.add_run('.')
                        run.font.name = 'Hanken Grotesk'
                        run.font.size = Pt(10)
                        run.font.weight = 300
                        run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    print(f"❌ DEBUG: No URL found for bullet point {bullet_point_count + 1}")

                bullet_point_count += 1

            # Regular paragraphs
            else:
                run = p.add_run(para.strip())
                run.font.name = 'Hanken Grotesk'
                run.font.size = Pt(10)
                run.font.weight = 300
                run.font.color.rgb = RGBColor(0, 0, 0)

    print(f"🏁 DEBUG: Finished processing. Total bullet points: {bullet_point_count}")

    # KEY CHANGES:
    # 1. More robust "Relevant Comments" detection using 'in' instead of '=='
    # 2. Multiple bullet character support (•, -, *, plus any line with REF-)
    # 3. Better debug output to see exactly what's happening
    # 4. More detailed error reporting

    filename = f'new_bios/{mp_name}_biography.docx'
    doc.save(filename)
    return filename



def main():
    # Check for API key
    if not os.getenv('ANTHROPIC_API_KEY'):
        print("Please set your ANTHROPIC_API_KEY first!")
        return

    # Get MP name
    mp_name = input("Enter MP name: ").strip()

    # Check if biography already exists
    if os.path.exists(f'new_bios/{mp_name}_biography.docx'):
        print(f"Biography already exists for {mp_name}")
        return

    # Get input file path
    input_path = input("Enter path to PDF file: ").strip()
    if not os.path.exists(input_path):
        print("File not found!")
        return

    try:
        # Read examples
        print("Reading example biographies...")
        examples = read_example_bios()

        # Read PDF input
        print("Reading PDF file...")
        input_content = read_pdf(input_path)
        has_pdf = input_content is not None


        # Get MP ID and API data
        mp_id = get_mp_id(mp_name)
        verified_positions = get_verified_positions(mp_id) if mp_id else None

        # Generate biography with verified positions
        biography = generate_biography(mp_name, input_content, examples, verified_positions)
        # Get Wikipedia data
        print("Fetching Wikipedia data...")
        wiki_data = get_wiki_data(mp_name)
        has_wiki_data = wiki_data is not None

        if input_content is None:
            print("Error reading PDF file")
            return

        # Generate biography
        print("Generating biography...")
        biography = generate_biography(mp_name, input_content, examples)

        # Get Wikipedia URL if data exists
        wiki_url = get_wiki_url(mp_name) if has_wiki_data else None

        # Save biography
        print("Saving biography...")
        saved_path = save_biography(mp_name, biography,
                                has_pdf=has_pdf,
                                has_api_data=has_api_data,
                                has_wiki_data=has_wiki_data,
                                wiki_url=wiki_url)
        print(f"Biography saved to {saved_path}")

    except Exception as e:
        print(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()
